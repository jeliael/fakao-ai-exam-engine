# -*- coding: utf-8 -*-
"""
generate_knowledge_card_docx.py — 知识卡片 Word 文档生成脚本

功能：
  读取 card.json，生成单个考点的知识卡片 Word 文档。
  格式规范：
    - 标题：宋体 16pt 居中加粗
    - 小标题：黑体 14pt 加粗
    - 正文：仿宋 小四(12pt)、1.5倍行距、首行缩进2字符
    - 页边距：上3.7cm / 下3.5cm / 左2.8cm / 右2.6cm

输入：
  --card  card.json 文件路径
输出：
  ../output/知识卡片-XX考点-YYYYMMDD.docx

用法：
  python generate_knowledge_card_docx.py --card card.json
  python generate_knowledge_card_docx.py --card card.json --output custom.docx

card.json 格式:
  {
    "subject": "刑法",
    "point": "共同犯罪",
    "concept": "概念定义...",
    "legal_basis": [
      {"law": "刑法", "article": "第25条", "text": "法条内容..."},
      ...
    ],
    "elements": ["构成要件1", "构成要件2", ...],
    "key_points": ["要点1", "要点2", ...],
    "common_confusion": [
      {"item": "易混点1", "distinction": "区分..."},
      ...
    ],
    "exam_tips": "应试技巧...",
    "related_points": ["关联考点1", "关联考点2"],
    "examples": [
      {"scenario": "案例...", "analysis": "分析...", "conclusion": "结论..."},
      ...
    ]
  }
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[错误] 缺少 python-docx 库，请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# 路径常量
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "..", "output")


# ---------------------------------------------------------------------------
# 数据加载
# ---------------------------------------------------------------------------
def load_json(file_path, label):
    """加载 JSON 文件。"""
    if not os.path.exists(file_path):
        print(f"[错误] {label}文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"[错误] {label}文件 JSON 解析失败: {e}", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Word 文档格式工具
# ---------------------------------------------------------------------------
def set_run_font(run, font_name="仿宋", size=12, bold=False, color=None):
    """设置 run 的字体（中文+西文）。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    """添加主标题：宋体 16pt 居中加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_name="宋体", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(6)


def add_subtitle(doc, text):
    """添加小标题：黑体 14pt 加粗。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="黑体", size=14, bold=True)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(12)


def add_body(doc, text, indent=True, bold=False, color=None):
    """添加正文：仿宋 12pt(小四)、1.5倍行距、首行缩进2字符。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="仿宋", size=12, bold=bold, color=color)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.space_after = Pt(3)


def add_bullet(doc, text, bold_prefix=None):
    """添加项目符号段落，可选加粗前缀。"""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        set_run_font(run1, font_name="仿宋", size=12, bold=True)
        run2 = p.add_run(text)
        set_run_font(run2, font_name="仿宋", size=12, bold=False)
    else:
        run = p.add_run(text)
        set_run_font(run, font_name="仿宋", size=12, bold=False)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, text, bold_prefix=None):
    """添加编号段落。"""
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        set_run_font(run1, font_name="仿宋", size=12, bold=True)
        run2 = p.add_run(text)
        set_run_font(run2, font_name="仿宋", size=12, bold=False)
    else:
        run = p.add_run(text)
        set_run_font(run, font_name="仿宋", size=12, bold=False)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows):
    """添加表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, font_name="黑体", size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, font_name="仿宋", size=11, bold=False)

    doc.add_paragraph()


def add_divider(doc):
    """添加分隔线。"""
    p = doc.add_paragraph()
    run = p.add_run("─" * 30)
    set_run_font(run, font_name="仿宋", size=10, bold=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)


def setup_page(doc):
    """设置页面：页边距 上3.7/下3.5/左2.8/右2.6 cm。"""
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


# ---------------------------------------------------------------------------
# 文档生成
# ---------------------------------------------------------------------------
def generate_knowledge_card(card_data, output_path):
    """生成知识卡片 Word 文档。"""
    doc = Document()
    setup_page(doc)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    date_str = datetime.now().strftime("%Y年%m月%d日")
    subject = card_data.get("subject", "未知科目")
    point = card_data.get("point", "未知考点")

    # === 标题 ===
    add_title(doc, f"知识卡片：{point}")
    add_body(doc, f"科目：{subject}　　考点：{point}　　日期：{date_str}", indent=False)
    add_divider(doc)

    # === 概念定义 ===
    concept = card_data.get("concept", "")
    if concept:
        add_subtitle(doc, "一、概念定义")
        add_body(doc, concept)

    # === 法条依据 ===
    legal_basis = card_data.get("legal_basis", [])
    if legal_basis:
        add_subtitle(doc, "二、法条依据")
        for lb in legal_basis:
            law_name = lb.get("law", "")
            article = lb.get("article", "")
            text = lb.get("text", "")
            add_body(doc, f"《{law_name}》{article}", indent=False, bold=True)
            if text:
                add_body(doc, text)
            add_body(doc, "")

    # === 构成要件 ===
    elements = card_data.get("elements", [])
    if elements:
        add_subtitle(doc, "三、构成要件")
        for i, element in enumerate(elements, 1):
            if isinstance(element, dict):
                add_numbered(doc, element.get("text", ""), bold_prefix=f"{element.get('name', '')}：")
            else:
                add_numbered(doc, element)

    # === 核心要点 ===
    key_points = card_data.get("key_points", [])
    if key_points:
        add_subtitle(doc, "四、核心要点")
        for kp in key_points:
            if isinstance(kp, dict):
                add_bullet(doc, kp.get("text", ""), bold_prefix=f"{kp.get('name', '')}：")
            else:
                add_bullet(doc, kp)

    # === 易混点辨析 ===
    common_confusion = card_data.get("common_confusion", [])
    if common_confusion:
        add_subtitle(doc, "五、易混点辨析")
        headers = ["易混项", "区分要点"]
        rows = []
        for cc in common_confusion:
            if isinstance(cc, dict):
                rows.append([cc.get("item", ""), cc.get("distinction", "")])
            else:
                rows.append([cc, ""])
        add_table(doc, headers, rows)

    # === 典型案例 ===
    examples = card_data.get("examples", [])
    if examples:
        add_subtitle(doc, "六、典型案例")
        for i, example in enumerate(examples, 1):
            add_body(doc, f"案例 {i}：", indent=False, bold=True)
            scenario = example.get("scenario", "")
            if scenario:
                add_body(doc, f"案情：{scenario}")
            analysis = example.get("analysis", "")
            if analysis:
                add_body(doc, f"分析：{analysis}")
            conclusion = example.get("conclusion", "")
            if conclusion:
                add_body(doc, f"结论：{conclusion}", bold=True)
            add_body(doc, "")

    # === 应试技巧 ===
    exam_tips = card_data.get("exam_tips", "")
    if exam_tips:
        add_subtitle(doc, "七、应试技巧")
        add_body(doc, exam_tips, color=(0, 0, 180))

    # === 关联考点 ===
    related_points = card_data.get("related_points", [])
    if related_points:
        add_subtitle(doc, "八、关联考点")
        for rp in related_points:
            add_bullet(doc, rp)

    # === 备注 ===
    add_subtitle(doc, "九、备注")
    add_body(doc, "1. 本知识卡片基于法考大纲和历年真题整理，仅供学习参考。")
    add_body(doc, "2. 法条内容以国家法律法规数据库(flk.npc.gov.cn)公布的现行有效版本为准。")
    add_body(doc, "3. 建议结合真题练习巩固理解，注意易混点的区分。")
    add_body(doc, f"\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 保存
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[完成] 知识卡片已保存至: {output_path}")


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="知识卡片 Word 文档生成 — 读取 card.json 生成单个考点的知识卡片",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_knowledge_card_docx.py --card card.json
  python generate_knowledge_card_docx.py --card card.json --output custom.docx
        """,
    )
    parser.add_argument("--card", required=True, type=str, help="知识卡片 JSON 文件路径")
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="输出文件路径（默认: ../output/知识卡片-XX考点-YYYYMMDD.docx）",
    )

    args = parser.parse_args()

    card_data = load_json(args.card, "知识卡片")

    if args.output:
        output_path = args.output
    else:
        date_str = datetime.now().strftime("%Y%m%d")
        point = card_data.get("point", "未知考点")
        output_path = os.path.join(OUTPUT_DIR, f"知识卡片-{point}-{date_str}.docx")

    generate_knowledge_card(card_data, output_path)


if __name__ == "__main__":
    main()
