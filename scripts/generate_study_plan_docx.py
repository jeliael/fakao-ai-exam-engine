# -*- coding: utf-8 -*-
"""
generate_study_plan_docx.py — 学习计划 Word 文档生成脚本

功能：
  读取 plan.json，生成排版规范的学习计划 Word 文档。
  格式规范：
    - 标题：宋体 16pt 居中加粗
    - 小标题：黑体 14pt 加粗
    - 正文：仿宋 小四(12pt)、1.5倍行距、首行缩进2字符
    - 页边距：上3.7cm / 下3.5cm / 左2.8cm / 右2.6cm

输入：
  --plan  plan.json 文件路径
输出：
  ../output/法考学习计划-YYYYMMDD.docx

用法：
  python generate_study_plan_docx.py --plan plan.json
  python generate_study_plan_docx.py --plan plan.json --output custom.docx

plan.json 格式:
  {
    "user": {"name": "张三", "level": "零基础", "target_date": "2026-09-01", "daily_hours": 4},
    "phases": [
      {"name": "基础阶段", "duration": "2025-09 ~ 2026-03", "goal": "...", "tasks": ["..."]},
      ...
    ],
    "daily_schedule": [
      {"time": "09:00-11:00", "task": "刑法精讲", "detail": "..."},
      ...
    ],
    "priority_points": [
      {"subject": "刑法", "point": "共同犯罪", "priority": "高", "reason": "..."},
      ...
    ]
  }
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    print("[错误] 缺少 python-docx 库，请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# 路径常量
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "..", "output")


# ---------------------------------------------------------------------------
# 数据加载
# ---------------------------------------------------------------------------
def load_json(file_path, label):
    """加载 JSON 文件。"""
    if not os.path.exists(file_path):
        print(f"[错误] {label}文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"[错误] {label}文件 JSON 解析失败: {e}", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Word 文档格式工具
# ---------------------------------------------------------------------------
def set_cell_font(run, font_name="仿宋", size=12, bold=False):
    """设置 run 的字体（中文+西文）。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_title(doc, text):
    """添加主标题：宋体 16pt 居中加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_cell_font(run, font_name="宋体", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(6)


def add_subtitle(doc, text):
    """添加小标题：黑体 14pt 加粗。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cell_font(run, font_name="黑体", size=14, bold=True)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(12)


def add_body(doc, text, indent=True):
    """添加正文：仿宋 12pt(小四)、1.5倍行距、首行缩进2字符。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cell_font(run, font_name="仿宋", size=12, bold=False)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)  # 约2字符
    p.paragraph_format.space_after = Pt(3)


def add_bullet(doc, text):
    """添加项目符号段落。"""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_cell_font(run, font_name="仿宋", size=12, bold=False)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows):
    """添加表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_cell_font(run, font_name="黑体", size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_cell_font(run, font_name="仿宋", size=11, bold=False)

    # 表后空行
    doc.add_paragraph()


def setup_page(doc):
    """设置页面：页边距 上3.7/下3.5/左2.8/右2.6 cm。"""
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


# ---------------------------------------------------------------------------
# 文档生成
# ---------------------------------------------------------------------------
def generate_study_plan(plan_data, output_path):
    """生成学习计划 Word 文档。"""
    doc = Document()
    setup_page(doc)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    date_str = datetime.now().strftime("%Y年%m月%d日")
    user = plan_data.get("user", {})

    # === 标题 ===
    add_title(doc, "法考学习计划")

    # === 用户信息 ===
    add_subtitle(doc, "一、学员信息")
    info_lines = [
        f"姓　　名：{user.get('name', '未知')}",
        f"基础水平：{user.get('level', '未知')}",
        f"目标考试日期：{user.get('target_date', '未知')}",
        f"每日学习时长：{user.get('daily_hours', '未知')} 小时",
        f"计划生成日期：{date_str}",
    ]
    for line in info_lines:
        add_body(doc, line)

    # === 阶段规划 ===
    phases = plan_data.get("phases", [])
    if phases:
        add_subtitle(doc, "二、阶段规划")
        for i, phase in enumerate(phases, 1):
            add_body(doc, f"（{i}）{phase.get('name', '')}　{phase.get('duration', '')}", indent=False)
            if phase.get("goal"):
                add_body(doc, f"阶段目标：{phase['goal']}")
            tasks = phase.get("tasks", [])
            if tasks:
                add_body(doc, "主要任务：", indent=False)
                for task in tasks:
                    add_bullet(doc, task)

    # === 每日任务安排 ===
    daily_schedule = plan_data.get("daily_schedule", [])
    if daily_schedule:
        add_subtitle(doc, "三、每日任务安排")
        headers = ["时间段", "学习任务", "具体内容"]
        rows = []
        for item in daily_schedule:
            rows.append([
                item.get("time", ""),
                item.get("task", ""),
                item.get("detail", ""),
            ])
        add_table(doc, headers, rows)

    # === 考点优先级 ===
    priority_points = plan_data.get("priority_points", [])
    if priority_points:
        add_subtitle(doc, "四、考点优先级排序")
        headers = ["科目", "考点", "优先级", "优先原因"]
        rows = []
        for item in priority_points:
            rows.append([
                item.get("subject", ""),
                item.get("point", ""),
                item.get("priority", ""),
                item.get("reason", ""),
            ])
        add_table(doc, headers, rows)

    # === 备注 ===
    add_subtitle(doc, "五、备注")
    add_body(doc, "1. 本计划基于个人基础和目标日期自动生成，请根据实际进度动态调整。")
    add_body(doc, "2. 建议每周进行一次小结，每月进行一次模考检测。")
    add_body(doc, "3. 考前两个月进入冲刺阶段，以真题和模考为主。")
    add_body(doc, f"\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 保存
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[完成] 学习计划已保存至: {output_path}")


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="学习计划 Word 文档生成 — 读取 plan.json 生成排版规范的学习计划",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_study_plan_docx.py --plan plan.json
  python generate_study_plan_docx.py --plan plan.json --output custom.docx
        """,
    )
    parser.add_argument("--plan", required=True, type=str, help="学习计划 JSON 文件路径")
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="输出文件路径（默认: ../output/法考学习计划-YYYYMMDD.docx）",
    )

    args = parser.parse_args()

    plan_data = load_json(args.plan, "学习计划")

    if args.output:
        output_path = args.output
    else:
        date_str = datetime.now().strftime("%Y%m%d")
        output_path = os.path.join(OUTPUT_DIR, f"法考学习计划-{date_str}.docx")

    generate_study_plan(plan_data, output_path)


if __name__ == "__main__":
    main()
