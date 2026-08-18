# -*- coding: utf-8 -*-
"""
generate_essay_feedback_docx.py — 主观题批改报告 Word 文档生成脚本

功能：
  读取 feedback.json，生成排版规范的主观题批改报告 Word 文档。
  格式规范：
    - 标题：宋体 16pt 居中加粗
    - 小标题：黑体 14pt 加粗
    - 正文：仿宋 小四(12pt)、1.5倍行距、首行缩进2字符
    - 页边距：上3.7cm / 下3.5cm / 左2.8cm / 右2.6cm

输入：
  --feedback  feedback.json 文件路径
输出：
  ../output/主观题批改报告-YYYYMMDD.docx

用法：
  python generate_essay_feedback_docx.py --feedback feedback.json
  python generate_essay_feedback_docx.py --feedback feedback.json --output custom.docx

feedback.json 格式:
  {
    "meta": {
      "question_type": "案例分析题",
      "total_score": 25,
      "essay_length": 800,
      "paragraph_count": 5,
      "status": "completed"
    },
    "dimensions": {
      "conclusion":     {"name": "结论正确性", "weight": 0.30, "max_score": 25, "score": 20, "weighted_score": 24.0, "criteria": "...", "comments": "..."},
      "legal_citation": {"name": "法条引用",   "weight": 0.20, "max_score": 25, "score": 18, "weighted_score": 14.4, "criteria": "...", "comments": "..."},
      ...
    },
    "weighted_total": 78.5,
    "paragraph_comments": [
      {"paragraph_index": 1, "paragraph": "...", "comment": "..."},
      ...
    ],
    "improvement_suggestions": ["建议1", "建议2", ...],
    "overall_comment": "总体评价...",
    "reference_answer": "参考答案...",
    "key_points": ["要点1", "要点2", ...]
  }
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[错误] 缺少 python-docx 库，请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# 路径常量
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "..", "output")


# ---------------------------------------------------------------------------
# 数据加载
# ---------------------------------------------------------------------------
def load_json(file_path, label):
    """加载 JSON 文件。"""
    if not os.path.exists(file_path):
        print(f"[错误] {label}文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"[错误] {label}文件 JSON 解析失败: {e}", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Word 文档格式工具
# ---------------------------------------------------------------------------
def set_run_font(run, font_name="仿宋", size=12, bold=False, color=None):
    """设置 run 的字体（中文+西文）。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    """添加主标题：宋体 16pt 居中加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_name="宋体", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(6)


def add_subtitle(doc, text):
    """添加小标题：黑体 14pt 加粗。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="黑体", size=14, bold=True)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(12)


def add_body(doc, text, indent=True, bold=False, color=None):
    """添加正文：仿宋 12pt(小四)、1.5倍行距、首行缩进2字符。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="仿宋", size=12, bold=bold, color=color)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.space_after = Pt(3)


def add_bullet(doc, text):
    """添加项目符号段落。"""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, font_name="仿宋", size=12, bold=False)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, highlight_col=None, highlight_threshold=None):
    """添加表格，可选高亮某列低于阈值的行。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, font_name="黑体", size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            # 高亮低分项
            if highlight_col is not None and col_idx == highlight_col:
                try:
                    val = float(str(cell_text).replace("%", ""))
                    if highlight_threshold is not None and val < highlight_threshold:
                        set_run_font(run, font_name="仿宋", size=11, bold=False, color=(255, 0, 0))
                    else:
                        set_run_font(run, font_name="仿宋", size=11, bold=False)
                except (ValueError, TypeError):
                    set_run_font(run, font_name="仿宋", size=11, bold=False)
            else:
                set_run_font(run, font_name="仿宋", size=11, bold=False)

    doc.add_paragraph()


def setup_page(doc):
    """设置页面：页边距 上3.7/下3.5/左2.8/右2.6 cm。"""
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


# ---------------------------------------------------------------------------
# 文档生成
# ---------------------------------------------------------------------------
def generate_essay_feedback(feedback_data, output_path):
    """生成主观题批改报告 Word 文档。"""
    doc = Document()
    setup_page(doc)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    date_str = datetime.now().strftime("%Y年%m月%d日")
    meta = feedback_data.get("meta", {})
    dimensions = feedback_data.get("dimensions", {})
    weighted_total = feedback_data.get("weighted_total")
    status = meta.get("status", "pending")

    # === 标题 ===
    add_title(doc, "主观题批改报告")

    # === 基本信息 ===
    add_subtitle(doc, "一、基本信息")
    info_lines = [
        f"题目类型：{meta.get('question_type', '未知')}",
        f"满分分值：{meta.get('total_score', '未知')} 分",
        f"答案字数：{meta.get('essay_length', '未知')} 字",
        f"段落数量：{meta.get('paragraph_count', '未知')} 段",
        f"批改状态：{'已完成' if status == 'completed' else '待AI评分'}",
        f"报告日期：{date_str}",
    ]
    for line in info_lines:
        add_body(doc, line)

    # === 各维度评分 ===
    if dimensions:
        add_subtitle(doc, "二、各维度评分")

        # 总分
        if weighted_total is not None:
            add_body(doc, f"加权总分：{weighted_total} / 100", bold=True)
        else:
            add_body(doc, "加权总分：待AI评分后计算", bold=True)
        add_body(doc, "")

        # 维度表格
        headers = ["评分维度", "权重", "得分", "满分", "加权得分", "评语"]
        rows = []
        dim_iter = dimensions.items() if isinstance(dimensions, dict) else enumerate(dimensions)
        for dim_key, dim in dim_iter:
            score = dim.get("score")
            score_str = str(score) if score is not None else "待评分"
            weighted = dim.get("weighted_score")
            weighted_str = str(weighted) if weighted is not None else "—"
            rows.append([
                dim.get("name", dim_key),
                f"{dim.get('weight', 0)*100:.0f}%",
                score_str,
                dim.get("max_score", ""),
                weighted_str,
                dim.get("comments", "") or dim.get("feedback", ""),
            ])
            # truncate last cell if too long
            if len(rows[-1][-1]) > 50:
                rows[-1][-1] = rows[-1][-1][:50] + "..."
        add_table(doc, headers, rows, highlight_col=2, highlight_threshold=1)

        # 各维度详细评语
        add_subtitle(doc, "各维度详细评语")
        dim_iter2 = dimensions.items() if isinstance(dimensions, dict) else enumerate(dimensions)
        for dim_key, dim in dim_iter2:
            name = dim.get("name", dim_key)
            score = dim.get("score")
            score_str = f"{score}/{dim.get('max_score', '')}" if score is not None else "待评分"
            add_body(doc, f"【{name}】得分：{score_str}（权重 {dim.get('weight', 0)*100:.0f}%）", indent=False)
            comments = dim.get("comments", "") or dim.get("feedback", "")
            if comments:
                add_body(doc, comments)
            else:
                add_body(doc, "（暂无评语）")

    # === 逐段批注 ===
    paragraph_comments = feedback_data.get("paragraph_comments", [])
    if paragraph_comments:
        add_subtitle(doc, "三、逐段批注")
        for pc in paragraph_comments:
            idx = pc.get("paragraph_index", "?")
            para_text = pc.get("paragraph", "")
            comment = pc.get("comment", "")

            add_body(doc, f"第 {idx} 段：", indent=False, bold=True)
            add_body(doc, f"原文：{para_text}")
            if comment:
                add_body(doc, f"批注：{comment}", color=(200, 0, 0))
            else:
                add_body(doc, "批注：（暂无）")
            add_body(doc, "")

    # === 改进建议 ===
    suggestions = feedback_data.get("improvement_suggestions", [])
    if suggestions:
        add_subtitle(doc, "四、改进建议")
        for i, suggestion in enumerate(suggestions, 1):
            add_body(doc, f"{i}. {suggestion}")

    # === 总体评价 ===
    overall_comment = feedback_data.get("overall_comment", "")
    if overall_comment:
        add_subtitle(doc, "五、总体评价")
        add_body(doc, overall_comment)

    # === 参考答案 ===
    reference_answer = feedback_data.get("reference_answer", "")
    if reference_answer:
        add_subtitle(doc, "六、参考答案")
        add_body(doc, reference_answer)

    # === 关键要点 ===
    key_points = feedback_data.get("key_points", [])
    if key_points:
        add_subtitle(doc, "七、答题关键要点")
        for i, point in enumerate(key_points, 1):
            add_bullet(doc, point)

    # === 备注 ===
    add_subtitle(doc, "八、备注")
    add_body(doc, "1. 本报告由 AI 辅助生成，评分仅供参考，最终成绩以司法部公布为准。")
    add_body(doc, "2. 主观题评分具有一定主观性，建议结合人工批改综合判断。")
    add_body(doc, "3. 红色文字表示批注意见或需重点关注的内容。")
    add_body(doc, f"\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 保存
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[完成] 主观题批改报告已保存至: {output_path}")


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="主观题批改报告 Word 文档生成 — 读取 feedback.json 生成批改报告",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_essay_feedback_docx.py --feedback feedback.json
  python generate_essay_feedback_docx.py --feedback feedback.json --output custom.docx
        """,
    )
    parser.add_argument("--feedback", required=True, type=str, help="批改反馈 JSON 文件路径")
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="输出文件路径（默认: ../output/主观题批改报告-YYYYMMDD.docx）",
    )

    args = parser.parse_args()

    feedback_data = load_json(args.feedback, "批改反馈")

    if args.output:
        output_path = args.output
    else:
        date_str = datetime.now().strftime("%Y%m%d")
        output_path = os.path.join(OUTPUT_DIR, f"主观题批改报告-{date_str}.docx")

    generate_essay_feedback(feedback_data, output_path)


if __name__ == "__main__":
    main()
