# -*- coding: utf-8 -*-
"""
generate_exam_report_docx.py — 模考报告 Word 文档生成脚本

功能：
  读取 report.json，生成排版规范的模考报告 Word 文档。
  格式规范：
    - 标题：宋体 16pt 居中加粗
    - 小标题：黑体 14pt 加粗
    - 正文：仿宋 小四(12pt)、1.5倍行距、首行缩进2字符
    - 页边距：上3.7cm / 下3.5cm / 左2.8cm / 右2.6cm

输入：
  --report  report.json 文件路径
输出：
  ../output/模考报告-YYYYMMDD.docx

用法：
  python generate_exam_report_docx.py --report report.json
  python generate_exam_report_docx.py --report report.json --output custom.docx

report.json 格式:
  {
    "exam_info": {"name": "模考第1轮", "date": "2025-08-17", "duration": "180分钟"},
    "summary": {"total_score": 85, "max_score": 150, "score_rate": 56.7, "rank": "需努力"},
    "subjects": [
      {"subject": "刑法", "score": 20, "max_score": 30, "accuracy": 66.7, "level": "一般"},
      ...
    ],
    "point_analysis": [
      {"point": "刑法>共同犯罪", "correct": 2, "total": 3, "accuracy": 66.7, "level": "一般"},
      ...
    ],
    "wrong_questions": [
      {"id": "Q005", "subject": "民法", "point": "善意取得", "user_answer": "B",
       "correct_answer": "ACD", "analysis": "..."},
      ...
    ],
    "suggestions": ["建议1", "建议2", ...]
  }
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[错误] 缺少 python-docx 库，请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# 路径常量
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "..", "output")


# ---------------------------------------------------------------------------
# 数据加载
# ---------------------------------------------------------------------------
def load_json(file_path, label):
    """加载 JSON 文件。"""
    if not os.path.exists(file_path):
        print(f"[错误] {label}文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"[错误] {label}文件 JSON 解析失败: {e}", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Word 文档格式工具
# ---------------------------------------------------------------------------
def set_run_font(run, font_name="仿宋", size=12, bold=False):
    """设置 run 的字体（中文+西文）。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_title(doc, text):
    """添加主标题：宋体 16pt 居中加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_name="宋体", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(6)


def add_subtitle(doc, text):
    """添加小标题：黑体 14pt 加粗。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="黑体", size=14, bold=True)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(12)


def add_body(doc, text, indent=True):
    """添加正文：仿宋 12pt(小四)、1.5倍行距、首行缩进2字符。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="仿宋", size=12, bold=False)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.space_after = Pt(3)


def add_bullet(doc, text):
    """添加项目符号段落。"""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, font_name="仿宋", size=12, bold=False)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows):
    """添加表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, font_name="黑体", size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, font_name="仿宋", size=11, bold=False)

    doc.add_paragraph()


def setup_page(doc):
    """设置页面：页边距 上3.7/下3.5/左2.8/右2.6 cm。"""
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


# ---------------------------------------------------------------------------
# 文档生成
# ---------------------------------------------------------------------------
def generate_exam_report(report_data, output_path):
    """生成模考报告 Word 文档。"""
    doc = Document()
    setup_page(doc)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    date_str = datetime.now().strftime("%Y年%m月%d日")
    exam_info = report_data.get("exam_info", {})
    summary = report_data.get("summary", {})

    # === 标题 ===
    add_title(doc, "法考模考报告")

    # === 考试信息 ===
    add_subtitle(doc, "一、考试信息")
    info_lines = [
        f"考试名称：{exam_info.get('name', '未知')}",
        f"考试日期：{exam_info.get('date', '未知')}",
        f"考试时长：{exam_info.get('duration', '未知')}",
        f"报告生成：{date_str}",
    ]
    for line in info_lines:
        add_body(doc, line)

    # === 成绩总览 ===
    add_subtitle(doc, "二、成绩总览")
    total_score = summary.get("total_score", 0)
    max_score = summary.get("max_score", 0)
    score_rate = summary.get("score_rate", 0)
    rank = summary.get("rank", "未知")

    summary_lines = [
        f"总分：{total_score} / {max_score}",
        f"得分率：{score_rate}%",
        f"评级：{rank}",
    ]
    for line in summary_lines:
        add_body(doc, line)

    # 评级说明
    if score_rate >= 80:
        add_body(doc, "评级说明：得分率80%以上，通过把握较大，继续保持。")
    elif score_rate >= 60:
        add_body(doc, "评级说明：得分率60%-80%，基本达标，需针对薄弱环节加强。")
    else:
        add_body(doc, "评级说明：得分率60%以下，需全面加强复习，建议调整学习策略。")

    # === 各科目表现 ===
    subjects = report_data.get("subjects", [])
    if subjects:
        add_subtitle(doc, "三、各科目表现")
        headers = ["科目", "得分", "满分", "得分率", "掌握程度"]
        rows = []
        for subj in subjects:
            rows.append([
                subj.get("subject", ""),
                subj.get("score", 0),
                subj.get("max_score", 0),
                f"{subj.get('accuracy', 0)}%",
                subj.get("level", ""),
            ])
        add_table(doc, headers, rows)

    # === 考点分析 ===
    point_analysis = report_data.get("point_analysis", [])
    if point_analysis:
        add_subtitle(doc, "四、考点掌握分析")
        headers = ["考点", "正确数", "总数", "正确率", "掌握程度"]
        rows = []
        for pt in point_analysis:
            rows.append([
                pt.get("point", ""),
                pt.get("correct", 0),
                pt.get("total", 0),
                f"{pt.get('accuracy', 0)}%",
                pt.get("level", ""),
            ])
        add_table(doc, headers, rows)

    # 薄弱考点提示
    weak_points = [p for p in point_analysis if p.get("level") == "薄弱"]
    if weak_points:
        add_body(doc, "薄弱考点（需重点复习）：", indent=False)
        for pt in weak_points:
            add_bullet(doc, f"{pt['point']}（正确率 {pt['accuracy']}%）")

    # === 错题解析 ===
    wrong_questions = report_data.get("wrong_questions", [])
    if wrong_questions:
        add_subtitle(doc, "五、错题解析")
        for i, wq in enumerate(wrong_questions, 1):
            add_body(doc, f"题目 {i}　（{wq.get('subject', '')} / {wq.get('point', '')}）", indent=False)
            add_body(doc, f"你的答案：{wq.get('user_answer', '未作答')}")
            add_body(doc, f"正确答案：{wq.get('correct_answer', '')}")
            analysis = wq.get("analysis", "")
            if analysis:
                add_body(doc, f"解析：{analysis}")
            add_body(doc, "")

    # === 改进建议 ===
    suggestions = report_data.get("suggestions", [])
    if suggestions:
        add_subtitle(doc, "六、改进建议")
        for i, suggestion in enumerate(suggestions, 1):
            add_body(doc, f"{i}. {suggestion}")

    # === 备注 ===
    add_subtitle(doc, "七、备注")
    add_body(doc, "1. 本报告基于本次模考数据自动生成，仅供参考。")
    add_body(doc, "2. 建议结合错题解析，针对薄弱考点制定专项复习计划。")
    add_body(doc, "3. 法考客观题满分150分（每题1-2分），主观题满分180分（分值分布因年而异）。")
    add_body(doc, f"\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 保存
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[完成] 模考报告已保存至: {output_path}")


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="模考报告 Word 文档生成 — 读取 report.json 生成排版规范的模考报告",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_exam_report_docx.py --report report.json
  python generate_exam_report_docx.py --report report.json --output custom.docx
        """,
    )
    parser.add_argument("--report", required=True, type=str, help="模考报告 JSON 文件路径")
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="输出文件路径（默认: ../output/模考报告-YYYYMMDD.docx）",
    )

    args = parser.parse_args()

    report_data = load_json(args.report, "模考报告")

    if args.output:
        output_path = args.output
    else:
        date_str = datetime.now().strftime("%Y%m%d")
        output_path = os.path.join(OUTPUT_DIR, f"模考报告-{date_str}.docx")

    generate_exam_report(report_data, output_path)


if __name__ == "__main__":
    main()
